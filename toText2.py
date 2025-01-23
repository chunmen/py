import docx
import json

def word_table_to_txt(word_file, txt_file):
    """
    将 Word 表格所有数据按规则处理后，每项数据存成一行 JSON 字符串。
    每 44 项为一个循环，添加空项使其变为 48 项，然后按分组倒序合并第 13-24 项到第 1-12 项。
    """
    try:
        doc = docx.Document(word_file)
    except FileNotFoundError:
        print(f"错误：找不到文件 '{word_file}'。请检查文件路径是否正确。")
        return
    except docx.opc.exceptions.PackageNotFoundError:
        print(f"错误：文件 '{word_file}' 不是有效的 Word 文件。")
        return
    except Exception as e:
        print(f"发生未知错误：{e}")
        return

    table_data = []

    # 遍历 Word 文档中的所有表格
    for table in doc.tables:
        table_info = []
        # 遍历表格中的所有行
        for row in table.rows:
            row_info = []
            # 遍历行中的所有单元格
            for cell in row.cells:
                # 提取单元格的格式信息（粗体和斜体）
                cell_format = {
                    "bold": cell.paragraphs[0].runs[0].bold if cell.paragraphs and cell.paragraphs[0].runs else False,
                    "italic": cell.paragraphs[0].runs[0].italic if cell.paragraphs and cell.paragraphs[0].runs else False,
                }
                row_info.append({"text": cell.text, "format": cell_format})
            table_info.append(row_info)
        table_data.append(table_info)

    all_items = []
    # 将所有表格的数据展平到一个列表中
    for table in table_data:
        for row in table:
            all_items.extend(row)

    # 添加空项
    all_items_with_empty = []
    for i, item in enumerate(all_items):
        all_items_with_empty.append(item)
        if i >= 192:  # 从第 193 项开始添加空项
            relative_index = (i - 192) % 44 + 1  # 计算在 44 项循环中的相对索引
            if relative_index in (29, 34, 39, 44):
                all_items_with_empty.append({"text": "", "format": {"bold": False, "italic": False}})  # 添加空项

    merged_all_items = []
    # 按每 48 项进行循环处理
    for start_index in range(0, len(all_items_with_empty), 48):
        end_index = min(start_index + 48, len(all_items_with_empty))  # 确保不超出列表范围
        current_items = all_items_with_empty[start_index:end_index] #获取当前48项或不足48项的数据

        filtered_items = []
        # 按照规则过滤数据（奇偶项删除）
        for i, item in enumerate(current_items):
            index = i + 1  # 索引从 1 开始
            group = (index - 1) // 12  # 计算当前项属于哪个 12 个一组的区间（从 0 开始）
            position_in_group = (index - 1) % 12 + 1  # 计算当前项在组内的位置（从 1 开始）

            if group % 2 == 0:  # 如果是第 0, 2, 4, ... 个区间（偶数区间）
                if position_in_group <= 24 and position_in_group % 2 != 0: #前24项奇数项保留
                    filtered_items.append(item)
                elif position_in_group > 24 and position_in_group <=36: #25到36项全部保留
                  filtered_items.append(item)
            else:  # 如果是第 1, 3, 5, ... 个区间（奇数区间）
                if position_in_group <= 24 and position_in_group % 2 != 0:
                    filtered_items.append(item)
                elif position_in_group > 24 and position_in_group <=36:
                  filtered_items.append(item)

        merged_items_in_loop = []
        # 合并第 13-24 项到第 1-12 项
        for i, item in enumerate(filtered_items):
            index = i + 1
            position_in_group = index  # 使用正确的索引 index

            if position_in_group <= 12:
                merged_items_in_loop.append(item)  # 前 12 项直接添加
            elif 13 <= position_in_group <= 24:
                # 计算合并的目标位置（按组倒序映射）
                group_index = (position_in_group - 13) // 3  # 计算当前项在 13-24 中的组索引（0, 1, 2, 3）
                position_in_group_in_group= (position_in_group - 13) % 3  #计算当前项在组内的位置(0,1,2)
                target_position = (group_index * 3) + (2 - position_in_group_in_group) + 1 # 计算目标位置

                if len(merged_items_in_loop) > target_position - 1:
                    if merged_items_in_loop[target_position - 1]["text"] != "":
                        merged_items_in_loop[target_position - 1]["text"] += "\n" + item["text"]
                    else:
                        merged_items_in_loop[target_position - 1]["text"] = item["text"]
                else:
                    merged_items_in_loop.append(item)

        merged_all_items.extend(merged_items_in_loop) #将每次循环的结果添加到最终结果中

    try:
        with open(txt_file, 'w', encoding='utf-8') as f:
            for item in merged_all_items:
                json_str = json.dumps(item, ensure_ascii=False)
                f.write(json_str + '\n')
        print(f"表格数据已成功保存到 '{txt_file}'。")
    except Exception as e:
        print(f"保存文件时发生错误：{e}")

# 示例用法
path =r"C:\Users\Administrator\Documents\WPSDrive\204442525\WPS云盘\新学生\陈司哲AI文章\英语\\"
word_file =path + "卡片打印csz.docx"
txt_file = path + "output.txt"

word_table_to_txt(word_file, txt_file)